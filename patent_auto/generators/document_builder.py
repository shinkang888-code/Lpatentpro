# -*- coding: utf-8 -*-
"""
특허 Word 문서 조립기
- KIPO / USPTO 양식에 맞는 Word 파일 자동 생성
- 섹션 데이터 + 도면 이미지를 받아 완전한 .docx 생성
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ─── 공통 유틸 ────────────────────────────────────────────────────────────
def _run(p, text, font_name="Times New Roman", size=12, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = font_name
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    try:
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    return r

def _heading(doc, text, font_name="Times New Roman", size=12, bold=True, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, font_name, size, bold)
    return p

def _para(doc, text, font_name="Times New Roman", size=12, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _run(p, text, font_name, size)
    return p

def _code(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(size)
    return p

def _insert_figure(doc, img_path: str, caption: str, width_cm=14):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_path, width=Cm(width_cm))
    except Exception as e:
        p = doc.add_paragraph(f"[도면 삽입 오류: {img_path} — {e}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    _run(cap, caption, size=10, bold=True)


# ─── KIPO 한국 출원용 문서 빌더 ───────────────────────────────────────────
class KIPODocumentBuilder:
    """특허법 제42조 및 특허법 시행규칙 제21조 준수 KIPO 문서 생성"""

    FONT = "맑은 고딕"

    def build(self, sections: dict, drawing_files: list, output_path: str):
        doc = Document()
        meta = sections.get("meta", {})
        inv = sections.get("invention", {})
        font = self.FONT

        # 페이지 설정 (A4)
        sec = doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(3.0)
        sec.bottom_margin = Cm(2.5)

        # ── 표지 ──────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "특  허  명  세  서", font, 20, bold=True)
        doc.add_paragraph()

        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = "Table Grid"
        rows_data = [
            ("출원인", meta.get("applicant_ko", "")),
            ("발명의 명칭", inv.get("title_ko", "")),
            ("출원 관할", "대한민국 특허청 (KIPO)"),
            ("IPC 분류", ", ".join(meta.get("ipc_codes", []))),
            ("작성 기준", "특허법 제42조 및 특허법 시행규칙 제21조"),
        ]
        for i, (k, v) in enumerate(rows_data):
            row = tbl.rows[i]
            rk = row.cells[0].paragraphs[0].add_run(k)
            rk.font.name = font; rk.font.size = Pt(10); rk.font.bold = True
            rk._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            rv = row.cells[1].paragraphs[0].add_run(v)
            rv.font.name = font; rv.font.size = Pt(10)
            rv._element.rPr.rFonts.set(qn('w:eastAsia'), font)

        doc.add_page_break()

        # ── 명세서 헤더 ─────────────────────────────────────────────────
        _heading(doc, "【명  세  서】", font, 14, center=True)
        _heading(doc, "【발명의 명칭】", font, 11)
        _para(doc, inv.get("title_ko", ""), font, 10)
        _para(doc, f"【IPC 분류】: {', '.join(meta.get('ipc_codes', []))}", font, 9)

        # ── 기술분야 ────────────────────────────────────────────────────
        _heading(doc, "【기술분야】", font, 11)
        for line in sections.get("technical_field", "").split("\n"):
            if line.strip():
                _para(doc, line, font, 10)

        # ── 배경기술 ────────────────────────────────────────────────────
        _heading(doc, "【발명의 배경이 되는 기술】", font, 11)
        bg = sections.get("background", [])
        for item in bg:
            if isinstance(item, dict):
                if item.get("para_num"):
                    _para(doc, item["para_num"], font, 9)
                _para(doc, item.get("text", ""), font, 10)
            else:
                _para(doc, str(item), font, 10)

        # ── 발명의 내용 ─────────────────────────────────────────────────
        _heading(doc, "【발명의 내용】", font, 11)
        summary = sections.get("summary", {})

        _heading(doc, "【해결하고자 하는 과제】", font, 10)
        _para(doc, "[0009]", font, 9)
        _para(doc, summary.get("problems", ""), font, 10)

        _heading(doc, "【과제의 해결 수단】", font, 10)
        _para(doc, "[0010]", font, 9)
        _para(doc, "상기 과제를 해결하기 위하여 본 발명은 다음의 구성을 포함한다:", font, 10)
        for line in summary.get("solutions", "").split("\n"):
            if line.strip():
                _para(doc, line, font, 10, indent=1)

        _heading(doc, "【발명의 효과】", font, 10)
        _para(doc, "[0011]", font, 9)
        for line in summary.get("effects", "").split("\n"):
            if line.strip():
                _para(doc, line, font, 10, indent=1)

        # ── 도면의 간단한 설명 ───────────────────────────────────────────
        _heading(doc, "【도면의 간단한 설명】", font, 11)
        _para(doc, "[0012]", font, 9)
        for fig in sections.get("brief_drawings", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _run(p, f"{fig['fig_no']}  ", font, 10, bold=True)
            _run(p, f"은 {fig['description']}", font, 10)

        # ── 발명의 구체적인 내용 ──────────────────────────────────────────
        _heading(doc, "【발명을 실시하기 위한 구체적인 내용】", font, 11)
        para_count = 13
        for item in sections.get("detailed", []):
            if item.get("section_title"):
                _heading(doc, f"§ {item['section_title']}", font, 10)
            _para(doc, item.get("para_num", f"[{para_count:04d}]"), font, 9)
            _para(doc, item.get("text", ""), font, 10)
            for algo in item.get("algorithms", []):
                _para(doc, f"[의사코드 — {algo.get('name', '')}]", font, 9)
                _code(doc, "\n".join(algo.get("pseudocode", [])))
            for formula in item.get("formulas", []):
                _para(doc, f"[{formula.get('label', '')}]  {formula.get('formula', '')}  — {formula.get('description', '')}", font, 10, indent=1)
            para_count += 1

        # ── 특허청구범위 ─────────────────────────────────────────────────
        doc.add_page_break()
        _heading(doc, "【특허청구범위】", font, 14, center=True)

        for claim in sections.get("claims", []):
            _heading(doc, f"【청구항 {claim['claim_no']}】", font, 11, bold=True)
            _para(doc, claim.get("text", ""), font, 10)
            doc.add_paragraph()

        # ── 요약서 ──────────────────────────────────────────────────────
        doc.add_page_break()
        _heading(doc, "【요  약  서】", font, 14, center=True)
        _heading(doc, "【요약】", font, 11)
        _para(doc, sections.get("abstract", ""), font, 10)
        _heading(doc, "【대표도】", font, 10)
        _para(doc, "도 1", font, 10)

        # ── 도면 ────────────────────────────────────────────────────────
        doc.add_page_break()
        _heading(doc, "【도  면】", font, 14, center=True)

        for i, fig_path in enumerate(drawing_files, 1):
            if os.path.exists(fig_path):
                caption = f"도 {i}"
                brief = next(
                    (f.get("description", "") for f in sections.get("brief_drawings", [])
                     if f.get("fig_no") == f"도 {i}"),
                    ""
                )
                _insert_figure(doc, fig_path, f"도 {i}. {brief}", width_cm=13)
            doc.add_paragraph()

        # 저장
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        print(f"  [OK] KIPO Word 저장: {output_path}")
        return output_path


# ─── USPTO 영문 문서 빌더 ────────────────────────────────────────────────
class USPTODocumentBuilder:
    """37 CFR 1.77 준수 USPTO 문서 생성"""

    FONT = "Times New Roman"

    def build(self, sections: dict, drawing_files: list, output_path: str):
        doc = Document()
        meta = sections.get("meta", {})
        inv = sections.get("invention", {})
        font = self.FONT

        # 페이지 설정 (8.5 x 11 inch)
        sec = doc.sections[0]
        sec.page_width = Cm(21.59)
        sec.page_height = Cm(27.94)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)

        # ── 표지 ─────────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "PATENT APPLICATION", font, 18, bold=True)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, "United States Patent and Trademark Office (USPTO)", font, 13)
        doc.add_paragraph()

        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = "Table Grid"
        rows_data = [
            ("Applicant:", meta.get("applicant_en", "")),
            ("Title:", inv.get("title_en", "")),
            ("Filing Authority:", "United States Patent and Trademark Office (USPTO)"),
            ("U.S. Classification:", ", ".join(meta.get("ipc_codes", []))),
            ("Prepared pursuant to:", "37 CFR 1.77 (Application Elements) and 37 CFR 1.84 (Drawings)"),
        ]
        for i, (k, v) in enumerate(rows_data):
            row = tbl.rows[i]
            rk = row.cells[0].paragraphs[0].add_run(k)
            rk.font.name = font; rk.font.size = Pt(10); rk.font.bold = True
            rv = row.cells[1].paragraphs[0].add_run(v)
            rv.font.name = font; rv.font.size = Pt(10)

        doc.add_page_break()

        # ── TITLE ────────────────────────────────────────────────────────
        _heading(doc, "TITLE OF THE INVENTION", font, 13, bold=True)
        _para(doc, f"This application is prepared in compliance with 37 CFR 1.77.", font, 10)
        _para(doc, inv.get("title_en", "").upper(), font, 12)

        # ── CROSS-REFERENCE ──────────────────────────────────────────────
        _heading(doc, "CROSS-REFERENCE TO RELATED APPLICATIONS", font, 13, bold=True)
        _para(doc, "[0001] This application claims priority to a Korean Patent Application filed with the Korean Intellectual Property Office (KIPO), the entire contents of which are incorporated herein by reference.", font, 12)

        # ── FIELD ────────────────────────────────────────────────────────
        _heading(doc, "FIELD OF THE INVENTION", font, 13, bold=True)
        field = sections.get("technical_field", "").replace("[0001]", "[0002]")
        _para(doc, field, font, 12)

        # ── BACKGROUND ───────────────────────────────────────────────────
        _heading(doc, "BACKGROUND OF THE INVENTION", font, 13, bold=True)
        bg = sections.get("background", [])
        para_num = 3
        for item in bg:
            if isinstance(item, dict):
                text = item.get("text", "")
                text = text.replace(f"[{para_num-1:04d}]", f"[{para_num:04d}]")
                _para(doc, f"[{para_num:04d}] {text}", font, 12)
                para_num += 1
            else:
                _para(doc, str(item), font, 12)

        # ── SUMMARY ──────────────────────────────────────────────────────
        _heading(doc, "SUMMARY OF THE INVENTION", font, 13, bold=True)
        summary = sections.get("summary", {})
        _para(doc, f"[{para_num:04d}] To solve the above technical problems, the present invention provides the following components:", font, 12)
        para_num += 1
        for line in summary.get("solutions", "").split("\n"):
            if line.strip():
                _para(doc, line, font, 12, indent=1.5)

        # ── BRIEF DESCRIPTION OF DRAWINGS ────────────────────────────────
        _heading(doc, "BRIEF DESCRIPTION OF THE DRAWINGS", font, 13, bold=True)
        _para(doc, f"[{para_num:04d}] The accompanying drawings, which are incorporated herein, illustrate embodiments of the invention.", font, 12)
        para_num += 1
        for fig in sections.get("brief_drawings", []):
            fig_label = fig.get("fig_no", "").replace("도", "FIG.")
            p = doc.add_paragraph()
            _run(p, f"{fig_label} ", font, 12, bold=True)
            _run(p, f"is {fig.get('description', '')}", font, 12)

        # ── DETAILED DESCRIPTION ─────────────────────────────────────────
        _heading(doc, "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS", font, 13, bold=True)
        for item in sections.get("detailed", []):
            if item.get("section_title"):
                _heading(doc, item["section_title"], font, 12, bold=True)
            _para(doc, f"[{para_num:04d}] {item.get('text', '')}", font, 12)
            para_num += 1
            for algo in item.get("algorithms", []):
                _para(doc, f"[Pseudocode — {algo.get('name', '')}]", font, 10, italic=True)
                _code(doc, "\n".join(algo.get("pseudocode", [])))
            for formula in item.get("formulas", []):
                _para(doc, f"[{formula.get('label', '')}]  {formula.get('formula', '')}  — {formula.get('description', '')}", font, 12, indent=1.5)

        # ── CLAIMS ───────────────────────────────────────────────────────
        doc.add_page_break()
        _heading(doc, "CLAIMS", font, 14, bold=True, center=True)

        for claim in sections.get("claims", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            _run(p, f"{claim['claim_no']}. ", font, 12, bold=True)
            _run(p, claim.get("text", ""), font, 12)
            doc.add_paragraph()

        # ── ABSTRACT ─────────────────────────────────────────────────────
        doc.add_page_break()
        _heading(doc, "ABSTRACT OF THE DISCLOSURE", font, 14, bold=True, center=True)
        _para(doc, sections.get("abstract", ""), font, 12)

        # ── DRAWINGS ────────────────────────────────────────────────────
        doc.add_page_break()
        _heading(doc, "DRAWINGS", font, 14, bold=True, center=True)

        for i, fig_path in enumerate(drawing_files, 1):
            if os.path.exists(fig_path):
                brief = next(
                    (f.get("description", "") for f in sections.get("brief_drawings", [])
                     if f.get("fig_no") == f"FIG. {i}"),
                    ""
                )
                _insert_figure(doc, fig_path, f"FIG. {i}  {brief}", width_cm=13)
            doc.add_paragraph()

        # 저장
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        print(f"  [OK] USPTO Word 저장: {output_path}")
        return output_path
