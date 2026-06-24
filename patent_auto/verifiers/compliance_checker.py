# -*- coding: utf-8 -*-
"""
특허 명세서 양식 자동 검증기
- KIPO 특허법 제42조 준수 검사
- USPTO 37 CFR 1.77 준수 검사
- 도면-명세서 일치성 검사
- 청구항 구조 검사
"""

import os
from pathlib import Path
from docx import Document


class ComplianceChecker:
    """특허 명세서 최종 검증"""

    def __init__(self):
        self.results = {"ok": [], "warn": [], "fail": []}

    def _check(self, label: str, condition: bool, level: str = "ok"):
        if condition:
            self.results["ok"].append(f"[OK] {label}")
        else:
            lvl = "fail" if level == "fail" else "warn"
            self.results[lvl].append(f"[{'FAIL' if lvl=='fail' else 'WARN'}] {label}")

    def check_kipo(self, docx_path: str) -> dict:
        """KIPO 특허법 제42조 준수 검사"""
        print("\n[검증] KIPO 양식 검사 중...")
        self.results = {"ok": [], "warn": [], "fail": []}

        if not os.path.exists(docx_path):
            self._check("KIPO 파일 존재", False, "fail")
            return self.results

        self._check("KIPO 파일 존재", True)
        sz = os.path.getsize(docx_path)
        self._check(f"파일 크기 적정 (현재 {sz//1024}KB)", sz > 50*1024)

        doc = Document(docx_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        # 필수 섹션
        kipo_required = {
            "【명세서】 또는 【명  세  서】": "명  세  서" in text or "명세서" in text,
            "【발명의 명칭】": "발명의 명칭" in text,
            "【기술분야】": "기술분야" in text,
            "【발명의 배경이 되는 기술】": "배경이 되는 기술" in text,
            "【발명의 내용】": "발명의 내용" in text,
            "【해결하고자 하는 과제】": "해결하고자 하는 과제" in text,
            "【과제의 해결 수단】": "과제의 해결 수단" in text,
            "【발명의 효과】": "발명의 효과" in text,
            "【도면의 간단한 설명】": "도면의 간단한 설명" in text,
            "【발명을 실시하기 위한 구체적인 내용】": "실시하기 위한 구체적인 내용" in text,
            "【특허청구범위】": "특허청구범위" in text,
            "【요약서】": "요약서" in text or "요  약  서" in text,
        }
        for name, cond in kipo_required.items():
            self._check(f"필수 섹션: {name}", cond, "fail" if not cond else "ok")

        # 단락번호 형식 [XXXX]
        for num in ["[0001]", "[0009]", "[0010]"]:
            self._check(f"단락번호 {num} 존재", num in text)

        # 청구항 구조
        if "청구항 1" in text or "【청구항 1】" in text:
            self._check("청구항 1 존재", True)
        else:
            self._check("청구항 1 존재", False, "fail")

        # 한국어 확인
        self._check("한국어 텍스트 (발명)", "발명" in text)

        # 도면 참조
        self._check("도면 참조 존재 (도 1)", "도 1" in text or "도1" in text)

        return self.results

    def check_uspto(self, docx_path: str) -> dict:
        """USPTO 37 CFR 1.77 준수 검사"""
        print("\n[검증] USPTO 양식 검사 중...")
        self.results = {"ok": [], "warn": [], "fail": []}

        if not os.path.exists(docx_path):
            self._check("USPTO 파일 존재", False, "fail")
            return self.results

        self._check("USPTO 파일 존재", True)
        sz = os.path.getsize(docx_path)
        self._check(f"파일 크기 적정 (현재 {sz//1024}KB)", sz > 50*1024)

        doc = Document(docx_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        # 37 CFR 1.77 필수 섹션
        uspto_required = {
            "TITLE OF THE INVENTION": "TITLE OF THE INVENTION" in text,
            "CROSS-REFERENCE TO RELATED APPLICATIONS": "CROSS-REFERENCE" in text,
            "FIELD OF THE INVENTION": "FIELD OF THE INVENTION" in text,
            "BACKGROUND OF THE INVENTION": "BACKGROUND OF THE INVENTION" in text,
            "SUMMARY OF THE INVENTION": "SUMMARY OF THE INVENTION" in text,
            "BRIEF DESCRIPTION OF THE DRAWINGS": "BRIEF DESCRIPTION" in text,
            "DETAILED DESCRIPTION": "DETAILED DESCRIPTION" in text,
            "CLAIMS": "CLAIMS" in text,
            "ABSTRACT": "ABSTRACT" in text,
            "37 CFR 준수 표기": "37 CFR" in text,
        }
        for name, cond in uspto_required.items():
            self._check(f"37CFR1.77 필수 섹션: {name}", cond, "fail" if not cond else "ok")

        # 청구항 번호 형식 "1."
        self._check("USPTO 청구항 1. 형식 존재", "1." in text)

        # 영어 텍스트
        self._check("영문 텍스트 (invention)", "invention" in text.lower())

        # FIG. 참조
        self._check("FIG. 1 도면 참조 존재", "FIG. 1" in text or "FIG.1" in text)

        return self.results

    def check_drawing_consistency(self, sections_ko: dict, sections_en: dict, 
                                   drawing_files_ko: list, drawing_files_en: list) -> dict:
        """도면-명세서 일치성 검사"""
        print("\n[검증] 도면-명세서 일치성 검사 중...")
        self.results = {"ok": [], "warn": [], "fail": []}

        # 도면 파일 존재 여부
        self._check(f"한국어 도면 생성됨 ({len(drawing_files_ko)}종)", len(drawing_files_ko) > 0)
        self._check(f"영문 도면 생성됨 ({len(drawing_files_en)}종)", len(drawing_files_en) > 0)

        for f in drawing_files_ko:
            self._check(f"KO 도면 파일 존재: {Path(f).name}", os.path.exists(f))
        for f in drawing_files_en:
            self._check(f"EN 도면 파일 존재: {Path(f).name}", os.path.exists(f))

        # 도면 수 일치
        brief_ko = sections_ko.get("brief_drawings", [])
        brief_en = sections_en.get("brief_drawings", [])
        self._check(f"한국/영문 도면 수 일치 ({len(brief_ko)}종)", len(brief_ko) == len(brief_en))

        # 청구항 수 일치
        claims_ko = sections_ko.get("claims", [])
        claims_en = sections_en.get("claims", [])
        self._check(f"한국/영문 청구항 수 일치 ({len(claims_ko)}개)", len(claims_ko) == len(claims_en))

        return self.results

    def print_report(self, title: str = "검증 결과"):
        """검증 결과 출력"""
        ok = len(self.results.get("ok", []))
        warn = len(self.results.get("warn", []))
        fail = len(self.results.get("fail", []))

        print(f"\n{'='*50}")
        print(f"[{title}]  통과: {ok}  경고: {warn}  실패: {fail}")
        print("=" * 50)
        for item in self.results.get("ok", []):
            print(f"  {item}")
        for item in self.results.get("warn", []):
            print(f"  {item}")
        for item in self.results.get("fail", []):
            print(f"  {item}")

        if fail == 0:
            print("\n  → 모든 검증 통과! 제출 준비 완료.")
        else:
            print(f"\n  → {fail}개 항목 수정 필요.")

        return fail == 0
